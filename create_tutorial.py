"""ETF 유튜브 대시보드 강의안 Word 문서 생성 (4회차 x 30분 구성)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return h

def add_para(text, bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 100, 0)
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    return p

def add_error_table(rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    headers[0].text = '증상'
    headers[1].text = '원인'
    headers[2].text = '해결법'
    for h in headers:
        for p in h.paragraphs:
            for r in p.runs:
                r.bold = True
    for symptom, cause, fix in rows:
        row = table.add_row().cells
        row[0].text = symptom
        row[1].text = cause
        row[2].text = fix
    doc.add_paragraph()

def add_checkpoint(items):
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(f'  {item}')
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.color.rgb = RGBColor(0, 128, 0)

# ════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('완전 초보자를 위한', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100, 100, 100))
add_para('Claude Code로 만드는\nETF 유튜브 데이터 대시보드', size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('컴퓨터를 처음 배우는 중학생도 혼자 완주할 수 있는 실습 강의안', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100, 100, 100))
doc.add_paragraph()
add_para('구성: 총 4회차 (회차당 30분) | 총 소요 시간: 약 2시간', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para('전제 조건: 인터넷 검색과 유튜브 시청이 가능하면 충분합니다!', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
doc.add_page_break()

# ════════════════════════════════════════════════════
# 시작하기 전에
# ════════════════════════════════════════════════════
add_heading_styled('시작하기 전에 - 이 강의안 사용법', level=1)
doc.add_paragraph()
add_para('이 강의안에서 자주 나오는 표시의 뜻:', bold=True)
doc.add_paragraph()
for icon, desc in [
    ('>> 중요한 개념 설명', ''),
    ('TIP - 알아두면 편한 것', ''),
    ('!! 막혔을 때 - 문제가 생겼을 때 해결법', ''),
    ('CHECK 체크포인트 - 여기까지 잘 따라왔는지 확인', ''),
]:
    p = doc.add_paragraph()
    p.add_run(icon).bold = True

doc.add_paragraph()

add_para('>> 막혔을 때 스크린샷으로 Claude에게 도움 요청하는 법', bold=True, size=13)
doc.add_paragraph('실습 중 에러가 나거나 막히면, 화면을 캡처해서 Claude Code에게 보여줄 수 있어요!')
doc.add_paragraph('이렇게 하면 Claude가 화면을 보고 정확하게 도와줍니다.')
doc.add_paragraph()
doc.add_paragraph('방법 1: 스크린샷 파일을 끌어다 놓기 (가장 쉬운 방법)')
doc.add_paragraph('  1) 키보드에서 "윈도우 키 + Shift + S"를 동시에 누르세요')
doc.add_paragraph('  2) 화면이 어두워지면, 캡처하고 싶은 부분을 마우스로 드래그하세요')
doc.add_paragraph('  3) 오른쪽 아래에 알림이 뜨면 클릭해서 저장하세요 (바탕화면에 저장 추천)')
doc.add_paragraph('  4) 저장된 스크린샷 파일을 Claude Code 터미널 창으로 끌어다 놓으세요 (드래그 앤 드롭)')
doc.add_paragraph('  5) 파일 경로가 자동으로 입력되면, 앞에 질문을 추가하세요:')
add_code_block('이 에러 화면 좀 봐줘 /Users/이름/Desktop/screenshot.png')
doc.add_paragraph()
doc.add_paragraph('방법 2: 클립보드에서 바로 붙여넣기')
doc.add_paragraph('  1) "윈도우 키 + Shift + S"로 화면 캡처 (클립보드에 자동 복사됨)')
doc.add_paragraph('  2) Claude Code에서 Ctrl+V로 붙여넣기')
doc.add_paragraph('  3) 이미지가 붙여넣어지면 질문과 함께 Enter!')
doc.add_paragraph()
doc.add_paragraph('방법 3: 파일 경로 직접 입력')
doc.add_paragraph('  바탕화면에 저장한 스크린샷의 경로를 직접 입력할 수도 있어요:')
add_code_block('이 화면에서 뭐가 잘못된 건지 알려줘 ~/Desktop/캡처.png')
doc.add_paragraph()
doc.add_paragraph('TIP: 에러 화면을 캡처할 때는 에러 메시지가 전부 보이게 넓게 캡처하세요!')
doc.add_paragraph('TIP: "이 에러 어떻게 해결해?", "이 화면이 맞아?" 같은 질문을 함께 적으면 더 정확한 답을 받을 수 있어요.')
doc.add_paragraph()

add_para('>> 용어 미리보기', bold=True, size=13)
terms = [
    ('터미널', '컴퓨터에게 글자로 명령을 내리는 채팅창. 카카오톡처럼 내가 글을 치면 컴퓨터가 답해줍니다.'),
    ('코드', '컴퓨터가 알아듣는 언어로 쓴 편지. 우리가 한국어를 쓰듯, 컴퓨터에게는 코드로 말합니다.'),
    ('Git', '코드의 타임머신. 내가 작업한 내용을 저장해두면, 언제든 과거로 돌아가거나 다른 사람과 공유할 수 있어요.'),
    ('Git Bash', 'Git과 함께 설치되는 특별한 터미널. 전 세계 개발자들이 쓰는 리눅스 스타일 명령어를 윈도우에서도 쓸 수 있게 해줘요.'),
    ('API', '앱과 앱 사이의 택배 시스템. 유튜브한테 "이 영상 정보 좀 줘"라고 부탁하는 우체통 같은 것.'),
    ('대시보드', '여러 정보를 한 화면에 보기 좋게 정리한 게시판. 자동차 계기판처럼 중요한 숫자를 한눈에 보여줍니다.'),
    ('ETF', '여러 회사의 주식을 한 바구니에 담아 파는 묶음 상품. 과자 종합선물세트 같은 것!'),
    ('PowerShell', '윈도우에 기본으로 있는 고급 터미널. 명령 프롬프트(cmd)의 업그레이드 버전이에요.'),
    ('명령 프롬프트(cmd)', '윈도우에 옛날부터 있던 가장 기본적인 터미널. 기능이 적어서 요즘은 잘 안 써요.'),
]
for term, desc in terms:
    p = doc.add_paragraph()
    run1 = p.add_run(f'  {term}: ')
    run1.bold = True
    p.add_run(desc)

doc.add_page_break()

# ════════════════════════════════════════════════════
# 1회차: 환경 설정 (터미널 + Git + Node.js/Python)
# ════════════════════════════════════════════════════
add_heading_styled('1회차: 개발 환경 설정 - 컴퓨터에게 말 걸 준비하기', level=1)
add_para('소요 시간: 30분', bold=True, color=(0, 100, 200))
add_para('이 회차를 마치면: 터미널, Git, Node.js, Python이 모두 설치되어 다음 회차에서 바로 코딩을 시작할 수 있어요!', bold=True)
doc.add_paragraph()

# ── 챕터 1: 윈도우 터미널 ──
add_heading_styled('챕터 1: 윈도우 터미널 설치 및 기본 사용법', level=2)
add_para('예상 시간: 7분', bold=True, color=(100, 100, 100))
doc.add_paragraph()

add_para('Step 1. Microsoft Store(마이크로소프트 스토어) 열기', bold=True, size=12)
doc.add_paragraph('키보드에서 윈도우 키(키보드 왼쪽 아래 창문 모양)를 한 번 누르세요.')
doc.add_paragraph('검색창이 뜨면 아래 글자를 그대로 입력하세요:')
add_code_block('Microsoft Store')
doc.add_paragraph('검색 결과에 파란 쇼핑백 모양 아이콘이 보이면 클릭하세요.')
doc.add_paragraph()

add_para('Step 2. Windows Terminal(윈도우 터미널) 검색 및 설치', bold=True, size=12)
doc.add_paragraph('Microsoft Store 앱이 열리면, 위쪽 검색창에 아래 글자를 입력하세요:')
add_code_block('Windows Terminal')
doc.add_paragraph('"Windows Terminal"이라고 적힌 앱이 보이면 클릭하고, 파란색 "설치" 또는 "다운로드" 버튼을 누르세요.')
doc.add_paragraph('TIP: 이미 "열기" 버튼이 보인다면? 이미 설치되어 있는 거예요! 바로 Step 3으로 넘어가세요.')
doc.add_paragraph()

add_para('Step 3. 터미널 처음 열어보기', bold=True, size=12)
doc.add_paragraph('설치가 끝나면 "열기" 버튼을 누르세요. (또는 윈도우 키를 누르고 "Terminal"을 검색해도 됩니다)')
doc.add_paragraph('까만(또는 파란) 창이 뜹니다. 이게 바로 터미널이에요!')
doc.add_paragraph('>> 무서워하지 마세요! 이건 컴퓨터랑 카톡하는 창이에요. 내가 글을 쓰면 컴퓨터가 대답합니다.')
doc.add_paragraph()
doc.add_paragraph('화면에 이런 비슷한 글자가 보일 거예요:')
add_code_block('PS C:\\Users\\여러분이름>')
doc.add_paragraph('이건 "나는 준비됐어, 명령을 내려줘!"라는 뜻이에요.')
doc.add_paragraph()

add_para('Step 4. 첫 번째 명령어 입력해보기', bold=True, size=12)
doc.add_paragraph('아래 명령어를 정확히 따라 입력하고 Enter(엔터) 키를 누르세요:')
add_code_block('echo "안녕, 나는 터미널이야!"')
doc.add_paragraph('화면에 이렇게 나와야 해요:')
add_code_block('안녕, 나는 터미널이야!')
doc.add_paragraph('축하해요! 방금 컴퓨터에게 첫 번째 명령을 내렸습니다!')
doc.add_paragraph()

add_para('>> 잠깐! 터미널이 여러 개라고요? - PowerShell, 명령 프롬프트(cmd), Git Bash 비교', bold=True, size=12)
doc.add_paragraph('윈도우에는 터미널(까만 창)이 여러 종류가 있어요. 처음에는 헷갈리지만, 비유로 쉽게 이해해봅시다!')
doc.add_paragraph()
doc.add_paragraph('  1) 명령 프롬프트(cmd): "기본 자전거"')
doc.add_paragraph('     - 윈도우에 처음부터 있는 가장 오래된 터미널')
doc.add_paragraph('     - 할 수 있는 것이 적고, 요즘 개발에는 잘 안 써요')
doc.add_paragraph('     - 화면에 "C:\\Users\\이름>" 이렇게 보여요')
doc.add_paragraph()
doc.add_paragraph('  2) PowerShell(파워셸): "전동 자전거"')
doc.add_paragraph('     - cmd를 업그레이드한 버전. 윈도우 10/11에 기본 설치되어 있어요')
doc.add_paragraph('     - cmd보다 많은 기능이 있지만, 개발자용 도구가 부족할 때가 있어요')
doc.add_paragraph('     - 화면에 "PS C:\\Users\\이름>" 이렇게 보여요 (앞에 PS가 붙어요)')
doc.add_paragraph()
doc.add_paragraph('  3) Git Bash: "전동 킥보드" (우리가 쓸 것!)')
doc.add_paragraph('     - Git을 설치하면 함께 오는 특별한 터미널')
doc.add_paragraph('     - 전 세계 개발자들이 쓰는 리눅스/맥 스타일 명령어를 윈도우에서도 쓸 수 있어요')
doc.add_paragraph('     - Claude Code가 가장 잘 동작하는 환경!')
doc.add_paragraph('     - 화면에 "이름@컴퓨터 MINGW64 ~" 이렇게 보이고, $ 표시로 명령을 기다려요')
doc.add_paragraph()
doc.add_paragraph('>> 결론: 이 강의에서는 Git Bash를 사용합니다! 다음 챕터에서 바로 설치할 거예요.')
doc.add_paragraph()

add_heading_styled('CHECK 체크포인트', level=3)
add_checkpoint([
    'CHECK 터미널 창이 열려 있나요?',
    'CHECK echo 명령어를 쳤을 때 내가 쓴 글이 그대로 화면에 나왔나요?',
])
doc.add_paragraph()

add_heading_styled('!! 막혔을 때', level=3)
add_error_table([
    ('Microsoft Store가 안 열려요', '윈도우 업데이트 필요', '윈도우 키 -> "업데이트 확인" 검색 -> 업데이트 실행 후 재시도'),
    ('"Windows Terminal"이 검색 안 돼요', '윈도우 버전이 낮음', '윈도우 키 -> "cmd" 검색 -> "명령 프롬프트"를 대신 사용'),
    ('글자를 쳤는데 아무 반응이 없어요', '커서가 터미널 밖에 있음', '터미널 까만 화면 안쪽을 마우스로 한 번 클릭하세요'),
])

# ── 챕터 2: Git 설치 ──
add_heading_styled('챕터 2: Git 설치 - 코드의 타임머신 준비하기', level=2)
add_para('예상 시간: 8분', bold=True, color=(100, 100, 100))
doc.add_paragraph()

add_para('>> Git을 왜 설치해야 하나요?', bold=True, size=12)
doc.add_paragraph('Git은 개발자에게 필수인 "코드 타임머신"이에요. Git이 필요한 이유 3가지:')
doc.add_paragraph()
doc.add_paragraph('  1) 실수 복구: 코드를 잘못 바꿨을 때 "아까 그 상태로 돌려줘!"가 가능해요.')
doc.add_paragraph('     비유: 게임의 "세이브 포인트"와 같아요. 중간중간 저장해두면 실패해도 다시 시작할 수 있죠!')
doc.add_paragraph()
doc.add_paragraph('  2) 협업: 여러 사람이 같은 프로젝트에서 동시에 작업할 수 있어요.')
doc.add_paragraph('     비유: 구글 문서에서 여러 명이 동시에 편집하는 것과 비슷해요.')
doc.add_paragraph()
doc.add_paragraph('  3) Git Bash: Git을 설치하면 "Git Bash"라는 특별한 터미널이 함께 설치돼요.')
doc.add_paragraph('     Claude Code는 이 Git Bash에서 가장 잘 동작합니다!')
doc.add_paragraph('     비유: 윈도우 기본 터미널이 "일반 자전거"라면, Git Bash는 "전동 킥보드" - 더 편하고 빨라요!')
doc.add_paragraph()

add_para('Step 1. Git 다운로드', bold=True, size=12)
doc.add_paragraph('브라우저에서 아래 주소로 이동하세요:')
add_code_block('https://git-scm.com/downloads/win')
doc.add_paragraph('화면에 "Click here to download" 또는 "64-bit Git for Windows Setup"을 클릭하세요.')
doc.add_paragraph()

add_para('Step 2. Git 설치하기', bold=True, size=12)
doc.add_paragraph('다운로드된 파일(Git-숫자-64-bit.exe)을 더블클릭하세요.')
doc.add_paragraph('"이 앱이 변경하도록 허용하시겠습니까?" -> "예"를 누르세요.')
doc.add_paragraph()
doc.add_paragraph('설치 화면이 나오면:')
doc.add_paragraph('  1) "Next" 버튼을 계속 누르세요 (약 8~10번)')
doc.add_paragraph('  2) 중간에 여러 옵션이 나오지만, 기본값 그대로 두고 "Next"만 누르면 돼요!')
doc.add_paragraph('  3) 마지막에 "Install" 버튼을 누르세요')
doc.add_paragraph('  4) 설치가 끝나면 "Finish" 버튼을 누르세요')
doc.add_paragraph()
doc.add_paragraph('TIP: 영어 옵션이 많이 나오지만 걱정 마세요! "Next -> Next -> ... -> Install -> Finish" 순서만 기억하면 됩니다.')
doc.add_paragraph()

add_para('Step 3. Git Bash 열어보기', bold=True, size=12)
doc.add_paragraph('바탕화면이나 시작 메뉴에서 "Git Bash"를 찾아 클릭하세요.')
doc.add_paragraph('(윈도우 키 -> "Git Bash" 검색 -> 클릭)')
doc.add_paragraph()
doc.add_paragraph('까만 창이 열리면서 이런 글자가 보여요:')
add_code_block('사용자이름@컴퓨터이름 MINGW64 ~\n$')
doc.add_paragraph('$ 표시가 "명령을 기다리고 있어요"라는 뜻이에요.')
doc.add_paragraph()
doc.add_paragraph('>> 앞으로 이 강의에서 "터미널을 열어주세요"라고 하면, 이 Git Bash를 열어주세요!')
doc.add_paragraph()

add_para('Step 4. Git 설치 확인', bold=True, size=12)
doc.add_paragraph('Git Bash에서 아래 명령어를 입력하고 Enter를 누르세요:')
add_code_block('git --version')
doc.add_paragraph('이런 결과가 나오면 성공:')
add_code_block('git version 2.47.1.windows.2')
doc.add_paragraph('(숫자는 다를 수 있어요. 숫자가 나오기만 하면 OK!)')
doc.add_paragraph()

add_para('Step 5. 프로젝트 폴더 만들기', bold=True, size=12)
doc.add_paragraph('앞으로 만들 파일들을 넣어둘 폴더(서랍장)를 만들어 봅시다.')
doc.add_paragraph('Git Bash에서 아래 명령어를 차례로 입력하세요:')
add_code_block('mkdir etf-dashboard\ncd etf-dashboard')
doc.add_paragraph('>> mkdir은 "make directory" = "폴더를 만들어줘"')
doc.add_paragraph('>> cd는 "change directory" = "이 폴더로 이동해줘"')
doc.add_paragraph()
doc.add_paragraph('화면이 이렇게 바뀌면 성공:')
add_code_block('사용자이름@컴퓨터이름 MINGW64 ~/etf-dashboard\n$')
doc.add_paragraph()

add_heading_styled('CHECK 체크포인트', level=3)
add_checkpoint([
    'CHECK Git Bash 창이 열려 있나요?',
    'CHECK git --version을 쳤을 때 숫자가 나왔나요?',
    'CHECK 화면에 "etf-dashboard"가 보이나요?',
])
doc.add_paragraph()

add_heading_styled('!! 막혔을 때', level=3)
add_error_table([
    ('Git Bash가 검색해도 안 나와요', 'Git 설치가 안 됐을 수 있어요', 'Git 설치 파일을 다시 다운로드하고 실행하세요'),
    ('"git"은(는) 인식할 수 없는 명령', '다른 터미널(PowerShell)을 쓰고 있어요', 'Git Bash를 열어주세요! (윈도우 키 -> "Git Bash" 검색)'),
    ('설치 중 옵션이 너무 많아서 헷갈려요', '기본값이면 충분해요', '모든 화면에서 아무것도 바꾸지 말고 "Next"만 누르세요'),
])

# ── 챕터 3: Node.js / Python ──
add_heading_styled('챕터 3: Node.js / Python 설치', level=2)
add_para('예상 시간: 15분', bold=True, color=(100, 100, 100))
doc.add_paragraph()

add_para('>> Node.js가 뭔가요?', bold=True, size=12)
doc.add_paragraph('Node.js(노드제이에스)는 프로그램을 실행시켜주는 "엔진"이에요.')
doc.add_paragraph('비유: 자동차(프로그램)가 달리려면 엔진(Node.js)이 필요한 것처럼, Claude Code를 돌리려면 Node.js가 필요해요.')
doc.add_paragraph()

add_para('Step 1. Node.js 다운로드 및 설치', bold=True, size=12)
doc.add_paragraph('브라우저에서 아래 주소로 이동하세요:')
add_code_block('https://nodejs.org')
doc.add_paragraph('"LTS"라고 적힌 큰 초록색 버튼을 클릭하세요.')
doc.add_paragraph('>> LTS = "Long Term Support" = "오래오래 안정적으로 쓸 수 있는 버전"')
doc.add_paragraph()
doc.add_paragraph('다운로드된 파일을 더블클릭하고:')
doc.add_paragraph('  "Next -> Next -> ... -> Install -> Finish" 순서로 진행하세요.')
doc.add_paragraph()

add_para('Step 2. Node.js 설치 확인', bold=True, size=12)
doc.add_paragraph('!! 중요: Git Bash를 껐다가 다시 열어야 합니다!')
doc.add_paragraph('Git Bash를 새로 열고 아래 명령어를 입력하세요:')
add_code_block('node --version')
doc.add_paragraph('이런 숫자가 나오면 성공:')
add_code_block('v22.13.1')
doc.add_paragraph()
doc.add_paragraph('이어서:')
add_code_block('npm --version')
doc.add_paragraph('>> npm = "Node Package Manager" = 앱스토어처럼 필요한 프로그램을 설치하는 도구')
doc.add_paragraph()

add_para('Step 3. Python 다운로드 및 설치', bold=True, size=12)
doc.add_paragraph('브라우저에서 아래 주소로 이동하세요:')
add_code_block('https://www.python.org/downloads/')
doc.add_paragraph('노란색 "Download Python 3.x.x" 버튼을 클릭하세요.')
doc.add_paragraph()
doc.add_paragraph('!! 가장 중요! 설치 화면 맨 아래에 "Add python.exe to PATH" 체크박스를 반드시 체크하세요!')
doc.add_paragraph('(체크 안 하면 나중에 큰 문제가 생겨요)')
doc.add_paragraph('그 다음 "Install Now"를 클릭하세요.')
doc.add_paragraph()

add_para('Step 4. Python 설치 확인', bold=True, size=12)
doc.add_paragraph('Git Bash를 껐다가 다시 열고:')
add_code_block('python --version')
doc.add_paragraph('이런 결과가 나오면 성공:')
add_code_block('Python 3.13.1')
doc.add_paragraph()

add_heading_styled('CHECK 체크포인트', level=3)
add_checkpoint([
    'CHECK node --version 결과에 v로 시작하는 숫자가 나왔나요?',
    'CHECK npm --version 결과에 숫자가 나왔나요?',
    'CHECK python --version 결과에 Python 3.x.x가 나왔나요?',
])
doc.add_paragraph()

add_heading_styled('!! 막혔을 때', level=3)
add_error_table([
    ('"node" 인식할 수 없는 명령', 'Git Bash를 새로 안 열었어요', 'Git Bash를 완전히 닫고 새로 여세요'),
    ('python 대신 Microsoft Store가 열려요', 'Windows 앱 별칭 문제', 'Git Bash에서는 보통 이 문제가 안 생겨요. 혹시 생기면 python3으로 시도'),
    ('설치 중 "관리자 권한" 에러', '권한 부족', '설치 파일 -> 마우스 우클릭 -> "관리자 권한으로 실행"'),
])

doc.add_paragraph()
add_para('1회차 완료! 수고하셨습니다! 다음 회차에서 Claude Code를 설치하고 AI와 대화해볼 거예요.', bold=True, size=12, color=(0, 128, 0))

doc.add_page_break()

# ════════════════════════════════════════════════════
# 2회차: Claude Code 설치 + API 키 발급
# ════════════════════════════════════════════════════
add_heading_styled('2회차: Claude Code 설치 + YouTube API 키 발급', level=1)
add_para('소요 시간: 30분', bold=True, color=(0, 100, 200))
add_para('이 회차를 마치면: AI 비서(Claude Code)와 터미널에서 대화할 수 있고, 유튜브에서 데이터를 가져올 수 있는 열쇠(API 키)도 준비돼요!', bold=True)
doc.add_paragraph()

# ── 챕터 4: Claude Code ──
add_heading_styled('챕터 4: Claude Code 설치 및 첫 실행', level=2)
add_para('예상 시간: 15분', bold=True, color=(100, 100, 100))
doc.add_paragraph()

add_para('>> Claude Code가 뭔가요?', bold=True, size=12)
doc.add_paragraph('Claude Code는 터미널 안에서 동작하는 AI 비서예요.')
doc.add_paragraph('비유: 카카오톡으로 똑똑한 친구에게 "이거 만들어줘"라고 부탁하는 것처럼, 터미널에서 Claude에게 프로그램을 만들어달라고 부탁할 수 있어요.')
doc.add_paragraph()

add_para('Step 1. Claude Code 설치', bold=True, size=12)
doc.add_paragraph('Git Bash를 열고 아래 명령어를 입력하세요:')
add_code_block('npm install -g @anthropic-ai/claude-code')
doc.add_paragraph('>> 이 명령어의 뜻:')
doc.add_paragraph('  npm install = "이 프로그램을 설치해줘"')
doc.add_paragraph('  -g = "어디서든 쓸 수 있게" (global의 줄임말)')
doc.add_paragraph('  @anthropic-ai/claude-code = "Claude Code라는 프로그램을"')
doc.add_paragraph()
doc.add_paragraph('설치에 1~3분 걸려요. 영어 글자가 쭉 올라가도 놀라지 마세요!')
doc.add_paragraph('마지막에 "added XX packages"가 보이면 성공!')
doc.add_paragraph()

add_para('Step 2. 프로젝트 폴더로 이동 후 Claude 실행', bold=True, size=12)
add_code_block('cd ~/etf-dashboard\nclaude')
doc.add_paragraph()
doc.add_paragraph('처음 실행하면 로그인이 필요합니다. 화면에 이런 메시지가 나타나요:')
add_code_block('? How would you like to authenticate? (Use arrow keys)\n> Anthropic Console (OAuth - recommended)\n  API Key')
doc.add_paragraph('키보드 화살표로 "Anthropic Console"이 선택된 상태에서 Enter를 누르세요.')
doc.add_paragraph('브라우저가 자동으로 열리면서 로그인 페이지가 나타납니다.')
doc.add_paragraph()

add_para('Step 3. 계정 만들기/로그인', bold=True, size=12)
doc.add_paragraph('브라우저에서:')
doc.add_paragraph('  - 계정이 있다면: 이메일과 비밀번호로 로그인')
doc.add_paragraph('  - 계정이 없다면: "Sign up"을 눌러 가입 (구글 계정으로도 가능)')
doc.add_paragraph('로그인 후 "Allow" 또는 "허용" 버튼을 누르세요.')
doc.add_paragraph()
doc.add_paragraph('터미널에 이런 화면이 보이면 성공:')
add_code_block('> Welcome to Claude Code!\n\n>')
doc.add_paragraph()

add_para('Step 4. Claude에게 첫 인사하기', bold=True, size=12)
doc.add_paragraph('> 표시 옆에 아래처럼 입력하고 Enter:')
add_code_block('안녕! 나는 ETF 대시보드를 만들고 싶어. 도와줄 수 있어?')
doc.add_paragraph('Claude가 한국어로 친절하게 대답해줄 거예요!')
doc.add_paragraph('TIP: Claude를 끝내려면 Ctrl+C 또는 /exit 입력')
doc.add_paragraph()

add_heading_styled('CHECK 체크포인트', level=3)
add_checkpoint([
    'CHECK claude 명령어로 "Welcome to Claude Code" 화면이 나왔나요?',
    'CHECK Claude에게 메시지를 보내고 답변을 받았나요?',
])
doc.add_paragraph()

add_heading_styled('!! 막혔을 때', level=3)
add_error_table([
    ('"claude" 인식할 수 없는 명령', 'npm 설치 실패 또는 경로 문제', 'Git Bash 닫고 다시 열기. 안 되면 npm install -g @anthropic-ai/claude-code 다시 실행'),
    ('로그인 페이지가 안 열려요', '브라우저 팝업 차단', '터미널에 나온 URL을 복사해서 브라우저 주소창에 직접 붙여넣기'),
    ('영어로 대답해요', '언어 자동 감지 실패', '"한국어로 대답해줘"라고 입력'),
])

# ── 챕터 5: YouTube API ──
add_heading_styled('챕터 5: YouTube Data API 키 발급 및 연결', level=2)
add_para('예상 시간: 15분', bold=True, color=(100, 100, 100))
doc.add_paragraph()

add_para('>> API 키가 뭔가요?', bold=True, size=12)
doc.add_paragraph('API 키는 "출입증"이에요. 놀이공원에 입장권이 필요하듯, 유튜브 데이터를 가져오려면 구글이 발급해주는 출입증(API 키)이 필요합니다.')
doc.add_paragraph('이 출입증은 무료예요! (하루 사용 횟수 제한만 있어요)')
doc.add_paragraph()

add_para('Step 1. Google Cloud Console 접속', bold=True, size=12)
doc.add_paragraph('브라우저에서 아래 주소로 이동하세요:')
add_code_block('https://console.cloud.google.com')
doc.add_paragraph('구글 계정으로 로그인하세요.')
doc.add_paragraph()

add_para('Step 2. 새 프로젝트 만들기', bold=True, size=12)
doc.add_paragraph('화면 왼쪽 위 "프로젝트 선택" (또는 "Select a project") 클릭')
doc.add_paragraph('"새 프로젝트" (또는 "New Project") 클릭')
doc.add_paragraph('프로젝트 이름:')
add_code_block('etf-youtube-dashboard')
doc.add_paragraph('"만들기" (또는 "Create") 클릭')
doc.add_paragraph()

add_para('Step 3. YouTube Data API 활성화', bold=True, size=12)
doc.add_paragraph('왼쪽 메뉴: "API 및 서비스" -> "라이브러리" (APIs & Services -> Library)')
doc.add_paragraph('검색창에 입력:')
add_code_block('YouTube Data API v3')
doc.add_paragraph('검색 결과에서 클릭 -> 파란색 "사용" (Enable) 버튼 클릭')
doc.add_paragraph()

add_para('Step 4. API 키 만들기', bold=True, size=12)
doc.add_paragraph('왼쪽 메뉴: "API 및 서비스" -> "사용자 인증 정보" (Credentials)')
doc.add_paragraph('"+ 사용자 인증 정보 만들기" -> "API 키" (Create Credentials -> API key)')
doc.add_paragraph()
doc.add_paragraph('API 키가 생성됩니다! 이런 모양의 긴 글자:')
add_code_block('AIzaSyD-xxxxxxxxxxxxxxxxxxxxxxxxx')
doc.add_paragraph('!! 매우 중요: 이 키를 복사해서 메모장에 저장하세요! 비밀번호처럼 남에게 보여주면 안 됩니다!')
doc.add_paragraph()

add_para('Step 5. API 키를 프로젝트에 저장', bold=True, size=12)
doc.add_paragraph('Git Bash에서:')
add_code_block('cd ~/etf-dashboard\necho "YOUTUBE_API_KEY=여러분의_API_키_여기에" > .env')
doc.add_paragraph('(여러분의_API_키_여기에 부분에 아까 복사한 키를 붙여넣으세요)')
doc.add_paragraph()
doc.add_paragraph('확인:')
add_code_block('cat .env')
doc.add_paragraph('결과:')
add_code_block('YOUTUBE_API_KEY=AIzaSyD-xxxxxxxxxxxxxxxxxxxxxxxxx')
doc.add_paragraph()

add_heading_styled('CHECK 체크포인트', level=3)
add_checkpoint([
    'CHECK YouTube Data API v3를 활성화(Enable) 했나요?',
    'CHECK API 키를 .env 파일에 저장했나요?',
    'CHECK cat .env 결과에 YOUTUBE_API_KEY=... 형태가 보이나요?',
])
doc.add_paragraph()

add_heading_styled('!! 막혔을 때', level=3)
add_error_table([
    ('"결제 계정 설정하세요" 메시지', '구글 결제 정보 요청', '무료 체험 등록 - 실제 과금은 안 돼요. 카드 등록은 본인 확인 용도'),
    ('API 키가 안 보여요', '팝업이 빨리 닫힘', '"사용자 인증 정보" 페이지에서 이미 만든 키 확인 가능'),
    ('"할당량 초과" 에러', '하루 사용량 초과', '24시간 후 재시도. 무료 할당량은 하루 10,000회로 충분'),
])

doc.add_paragraph()
add_para('2회차 완료! 수고하셨습니다! 다음 회차에서 드디어 데이터를 수집해볼 거예요.', bold=True, size=12, color=(0, 128, 0))

doc.add_page_break()

# ════════════════════════════════════════════════════
# 3회차: ETF 유튜브 데이터 수집 자동화
# ════════════════════════════════════════════════════
add_heading_styled('3회차: Claude Code로 ETF 유튜브 데이터 수집 자동화', level=1)
add_para('소요 시간: 30분', bold=True, color=(0, 100, 200))
add_para('이 회차를 마치면: Claude AI에게 부탁해서, 유튜브에서 ETF 관련 영상 정보를 자동으로 모아오는 프로그램이 완성돼요!', bold=True)
doc.add_paragraph()

add_para('>> 이번 회차의 핵심 아이디어', bold=True, size=12)
doc.add_paragraph('우리가 직접 코드를 쓰지 않아요! Claude AI에게 "이런 프로그램 만들어줘"라고 말하면, AI가 대신 코드를 작성해줍니다.')
doc.add_paragraph()

add_heading_styled('챕터 6: 데이터 수집 프로그램 만들기', level=2)
add_para('예상 시간: 15분', bold=True, color=(100, 100, 100))
doc.add_paragraph()

add_para('Step 1. Claude Code 실행', bold=True, size=12)
doc.add_paragraph('Git Bash에서:')
add_code_block('cd ~/etf-dashboard\nclaude')
doc.add_paragraph()

add_para('Step 2. Claude에게 데이터 수집 프로그램 만들어달라고 부탁하기', bold=True, size=12)
doc.add_paragraph('Claude의 > 프롬프트에 아래 내용을 통째로 복사해서 붙여넣으세요:')
doc.add_paragraph()
add_code_block(
    '.env 파일에 YOUTUBE_API_KEY가 저장되어 있어.\n'
    '이 키를 사용해서 유튜브에서 ETF 관련 영상을 검색하고\n'
    '데이터를 수집하는 Python 프로그램을 만들어줘.\n'
    '\n'
    '요구사항:\n'
    '1. "ETF 투자", "ETF 추천", "ETF 분석" 키워드로 검색\n'
    '2. 각 키워드당 최신 영상 10개씩 수집\n'
    '3. 수집할 정보: 영상 제목, 채널명, 조회수, 좋아요 수,\n'
    '   게시일, 영상 링크\n'
    '4. 결과를 etf_youtube_data.json 파일로 저장\n'
    '5. python-dotenv를 사용해서 .env에서 API 키 읽기\n'
    '6. 실행하면 수집 진행 상황을 보여주기'
)
doc.add_paragraph()
doc.add_paragraph('Enter를 누르면 Claude가 코드를 작성하기 시작합니다!')
doc.add_paragraph('TIP: Claude가 "파일을 만들어도 될까요?" 같은 질문을 하면 "y"를 입력하고 Enter (y = yes = 네)')
doc.add_paragraph()

add_para('Step 3. 만들어진 코드 확인', bold=True, size=12)
doc.add_paragraph('Claude가 작업을 마치면 /exit으로 나온 후:')
add_code_block('/exit\nls')
doc.add_paragraph('이런 파일들이 보여야 해요:')
add_code_block('.env\ncollect_data.py\nrequirements.txt')
doc.add_paragraph()

add_heading_styled('챕터 7: 데이터 수집 실행하기', level=2)
add_para('예상 시간: 15분', bold=True, color=(100, 100, 100))
doc.add_paragraph()

add_para('Step 4. 필요한 재료(라이브러리) 설치', bold=True, size=12)
add_code_block('pip install -r requirements.txt')
doc.add_paragraph('>> 이건 "requirements.txt에 적힌 재료들을 전부 설치해줘"라는 뜻이에요.')
doc.add_paragraph()

add_para('Step 5. 데이터 수집 프로그램 실행!', bold=True, size=12)
add_code_block('python collect_data.py')
doc.add_paragraph()
doc.add_paragraph('화면에 이런 진행 상황이 나타날 거예요:')
add_code_block(
    '"ETF 투자" 검색 중...\n'
    '  10개 영상 수집 완료\n'
    '"ETF 추천" 검색 중...\n'
    '  10개 영상 수집 완료\n'
    '"ETF 분석" 검색 중...\n'
    '  10개 영상 수집 완료\n'
    '\n'
    '총 30개 영상 데이터 수집 완료!\n'
    'etf_youtube_data.json에 저장했습니다.'
)
doc.add_paragraph()
doc.add_paragraph('축하해요! 유튜브에서 ETF 데이터를 자동으로 모아왔습니다!')
doc.add_paragraph()

add_para('Step 6. 수집된 데이터 확인해보기', bold=True, size=12)
add_code_block("python -c \"import json; data=json.load(open('etf_youtube_data.json',encoding='utf-8')); print(f'총 {len(data)}개 영상'); print(data[0]['title'])\"")
doc.add_paragraph('첫 번째 영상의 제목이 화면에 나올 거예요!')
doc.add_paragraph()

add_heading_styled('CHECK 체크포인트', level=3)
add_checkpoint([
    'CHECK collect_data.py 파일이 만들어졌나요?',
    'CHECK python collect_data.py 실행 시 "수집 완료" 메시지가 나왔나요?',
    'CHECK etf_youtube_data.json 파일이 만들어졌나요?',
])
doc.add_paragraph()

add_heading_styled('!! 막혔을 때', level=3)
add_error_table([
    ('ModuleNotFoundError', '재료가 설치 안 됨', 'pip install -r requirements.txt 다시 실행'),
    ('API key invalid / 403 에러', 'API 키 문제', '.env 파일의 키 재확인 + Google Cloud에서 API "사용" 상태 확인'),
    ('Claude가 다른 파일명을 썼어요', 'Claude가 때때로 다른 이름 선택', 'ls로 확인하고 Claude가 만든 파일명으로 실행'),
    ('인코딩 에러', '한글 처리 문제', 'Claude에게 "인코딩 에러 나요. utf-8로 수정해줘" 부탁'),
])

doc.add_paragraph()
add_para('3회차 완료! 수고하셨습니다! 마지막 회차에서 예쁜 대시보드를 만들어볼 거예요!', bold=True, size=12, color=(0, 128, 0))

doc.add_page_break()

# ════════════════════════════════════════════════════
# 4회차: 대시보드 제작 및 브라우저 확인
# ════════════════════════════════════════════════════
add_heading_styled('4회차: 대시보드 제작 및 브라우저에서 확인 - 프로젝트 완성!', level=1)
add_para('소요 시간: 30분', bold=True, color=(0, 100, 200))
add_para('이 회차를 마치면: 수집한 ETF 유튜브 데이터를 예쁜 웹 페이지(대시보드)로 만들어서 브라우저에서 확인! 프로젝트 완성!', bold=True)
doc.add_paragraph()

add_heading_styled('챕터 8: 대시보드 만들기', level=2)
add_para('예상 시간: 20분', bold=True, color=(100, 100, 100))
doc.add_paragraph()

add_para('>> 대시보드가 뭔가요?', bold=True, size=12)
doc.add_paragraph('자동차 운전석의 계기판(dashboard)처럼, 중요한 숫자와 그래프를 한 화면에 정리해서 보여주는 웹 페이지예요.')
doc.add_paragraph()

add_para('Step 1. Claude Code 실행', bold=True, size=12)
doc.add_paragraph('Git Bash에서:')
add_code_block('cd ~/etf-dashboard\nclaude')
doc.add_paragraph()

add_para('Step 2. Claude에게 대시보드 만들어달라고 부탁하기', bold=True, size=12)
doc.add_paragraph('Claude의 > 프롬프트에 아래 내용을 통째로 복사해서 붙여넣으세요:')
doc.add_paragraph()
add_code_block(
    'etf_youtube_data.json 파일에 유튜브 ETF 영상 데이터가 있어.\n'
    '이 데이터를 보여주는 예쁜 웹 대시보드를 HTML로 만들어줘.\n'
    '\n'
    '요구사항:\n'
    '1. 하나의 index.html 파일로 만들어줘 (서버 없이 바로 열 수 있게)\n'
    '2. 한국어로 작성\n'
    '3. 반드시 포함할 내용:\n'
    '   - 상단에 "ETF 유튜브 트렌드 대시보드" 제목\n'
    '   - 전체 영상 수, 총 조회수, 평균 좋아요 수 요약 카드\n'
    '   - 조회수 TOP 10 영상 막대 그래프 (Chart.js 사용)\n'
    '   - 채널별 영상 수 파이 차트\n'
    '   - 전체 영상 목록 테이블 (제목 클릭하면 유튜브로 이동)\n'
    '4. 디자인: 깔끔한 다크 테마, 카드형 레이아웃\n'
    '5. JSON 데이터를 HTML 파일 안에 직접 포함시켜줘\n'
    '6. 모바일에서도 보기 좋게 반응형으로 만들어줘'
)
doc.add_paragraph()
doc.add_paragraph('Enter를 누르면 Claude가 대시보드를 만들기 시작합니다!')
doc.add_paragraph('Claude가 파일을 만들어도 되는지 물으면 "y"를 입력하세요.')
doc.add_paragraph('코드 작성에 1~3분 걸릴 수 있어요. 기다려주세요!')
doc.add_paragraph()

add_para('Step 3. 대시보드 파일 확인', bold=True, size=12)
doc.add_paragraph('Claude가 작업을 마치면:')
add_code_block('/exit\nls')
doc.add_paragraph('index.html 파일이 보이면 성공!')
doc.add_paragraph()

add_heading_styled('챕터 9: 브라우저에서 대시보드 확인 + 꾸미기', level=2)
add_para('예상 시간: 10분', bold=True, color=(100, 100, 100))
doc.add_paragraph()

add_para('Step 4. 드디어! 브라우저에서 대시보드 열기', bold=True, size=12)
doc.add_paragraph('이 순간을 위해 여기까지 왔습니다!')
add_code_block('start index.html')
doc.add_paragraph()
doc.add_paragraph('인터넷 브라우저가 열리면서 여러분이 만든 대시보드가 나타납니다!')
doc.add_paragraph()
doc.add_paragraph('화면에 이런 것들이 보여요:')
doc.add_paragraph('  - 맨 위: "ETF 유튜브 트렌드 대시보드" 제목')
doc.add_paragraph('  - 요약 카드: 전체 영상 수, 총 조회수, 평균 좋아요 수')
doc.add_paragraph('  - 막대 그래프: 조회수 TOP 10 영상')
doc.add_paragraph('  - 파이 차트: 채널별 영상 수')
doc.add_paragraph('  - 영상 목록 테이블: 제목 클릭하면 유튜브로 이동!')
doc.add_paragraph()

add_para('Step 5. 대시보드 꾸미기 (도전 과제!)', bold=True, size=12)
doc.add_paragraph('마음에 안 드는 부분이 있다면? Claude에게 수정을 부탁하세요!')
add_code_block(
    'claude\n'
    '> index.html의 배경색을 좀 더 밝게 바꿔줘\n'
    '> 그래프 색상을 파란색 계열로 바꿔줘\n'
    '> 조회수 순 정렬 기능을 추가해줘'
)
doc.add_paragraph('수정 후 브라우저에서 Ctrl+R(새로고침)을 누르면 바뀐 화면을 볼 수 있어요!')
doc.add_paragraph()

add_heading_styled('CHECK 체크포인트', level=3)
add_checkpoint([
    'CHECK index.html 파일이 만들어졌나요?',
    'CHECK 브라우저에서 대시보드가 열렸나요?',
    'CHECK 그래프와 표가 화면에 보이나요?',
    'CHECK 영상 제목을 클릭하면 유튜브로 이동하나요?',
])
doc.add_paragraph()

add_heading_styled('!! 막혔을 때', level=3)
add_error_table([
    ('브라우저에 아무것도 안 보여요', 'HTML 파일 문제', 'Claude에게 "index.html이 안 열려요. 확인해줘" 부탁'),
    ('그래프가 안 보이고 빈 화면', '데이터 미포함', 'Claude에게 "그래프가 안 보여요. 데이터 다시 확인하고 HTML에 넣어줘" 부탁'),
    ('start 명령이 안 먹어요', '명령어 문제', 'etf-dashboard 폴더에서 index.html 파일을 직접 더블클릭'),
    ('글자가 깨져요', '인코딩 문제', 'Claude에게 "한글 깨져요. UTF-8 인코딩으로 수정해줘" 부탁'),
])

doc.add_paragraph()

# ── 챕터 10: GitHub ──
add_heading_styled('보너스 챕터: GitHub에 내 프로젝트 올리기', level=2)
add_para('예상 시간: 10분 (4회차 시간 내 포함)', bold=True, color=(100, 100, 100))
doc.add_paragraph()

add_para('>> GitHub(깃허브)가 뭔가요?', bold=True, size=12)
doc.add_paragraph('GitHub는 "코드의 인스타그램"이에요!')
doc.add_paragraph('  - 인스타그램에 사진을 올리듯, GitHub에 내 코드를 올려서 보관하고 공유할 수 있어요')
doc.add_paragraph('  - 전 세계 개발자들이 자기 프로젝트를 GitHub에 올려요')
doc.add_paragraph('  - 내 컴퓨터가 고장 나도, GitHub에 올려두면 코드가 안전하게 보관돼요')
doc.add_paragraph('  - 무료로 사용할 수 있어요!')
doc.add_paragraph()

add_para('>> Repository(저장소)가 뭔가요?', bold=True, size=12)
doc.add_paragraph('Repository(레포지토리, 줄여서 "레포")는 "프로젝트 폴더"예요.')
doc.add_paragraph('우리가 만든 etf-dashboard 폴더를 GitHub에 올리면, 그게 하나의 Repository가 됩니다.')
doc.add_paragraph('비유: 인스타그램의 "게시물 하나"가 Repository 하나라고 생각하면 돼요.')
doc.add_paragraph()

add_para('Step 1. GitHub 가입하기', bold=True, size=12)
doc.add_paragraph('브라우저에서 아래 주소로 이동하세요:')
add_code_block('https://github.com')
doc.add_paragraph()
doc.add_paragraph('오른쪽 위에 "Sign up" 버튼을 클릭하세요.')
doc.add_paragraph()
doc.add_paragraph('가입 화면이 나오면 차례대로 입력하세요:')
doc.add_paragraph('  1) "Enter your email" - 이메일 주소 입력 후 "Continue"')
doc.add_paragraph('  2) "Create a password" - 비밀번호 만들기 (8자 이상, 숫자 포함) 후 "Continue"')
doc.add_paragraph('  3) "Enter a username" - 사용자 이름 만들기 (영어, 숫자, - 만 사용 가능)')
doc.add_paragraph('     예시: my-first-code, student-kim, etf-learner')
doc.add_paragraph('     TIP: 이 이름이 내 GitHub 주소가 돼요! (github.com/내이름)')
doc.add_paragraph('  4) 이메일 수신 여부 - "n"을 입력해도 괜찮아요')
doc.add_paragraph('  5) 퍼즐 인증 - 화면의 퍼즐을 풀어주세요')
doc.add_paragraph('  6) "Create account" 클릭')
doc.add_paragraph('  7) 이메일로 온 인증 코드(6자리 숫자)를 입력하세요')
doc.add_paragraph()
doc.add_paragraph('가입 완료! 이제 GitHub 회원이에요!')
doc.add_paragraph()

add_para('Step 2. 새 Repository(저장소) 만들기', bold=True, size=12)
doc.add_paragraph('GitHub에 로그인한 상태에서:')
doc.add_paragraph()
doc.add_paragraph('  1) 오른쪽 위의 "+" 버튼을 클릭하세요')
doc.add_paragraph('  2) "New repository"를 클릭하세요')
doc.add_paragraph()
doc.add_paragraph('새 저장소 만들기 화면이 나와요:')
doc.add_paragraph()
doc.add_paragraph('  - "Repository name" (저장소 이름): 아래처럼 입력하세요')
add_code_block('etf-youtube-dashboard')
doc.add_paragraph('  - "Description" (설명, 선택사항): 아래처럼 입력하세요')
add_code_block('ETF YouTube Data Dashboard - My First Project')
doc.add_paragraph('  - "Public"을 선택하세요 (다른 사람도 볼 수 있게)')
doc.add_paragraph('    (비공개로 하고 싶다면 "Private" 선택)')
doc.add_paragraph('  - 나머지는 건드리지 마세요!')
doc.add_paragraph('  - 맨 아래 초록색 "Create repository" 버튼을 클릭!')
doc.add_paragraph()
doc.add_paragraph('빈 저장소가 만들어졌어요! 화면에 명령어들이 보일 거예요.')
doc.add_paragraph()

add_para('Step 3. 내 코드를 GitHub에 올리기', bold=True, size=12)
doc.add_paragraph('Git Bash를 열고 프로젝트 폴더로 이동하세요:')
add_code_block('cd ~/etf-dashboard')
doc.add_paragraph()
doc.add_paragraph('아래 명령어를 한 줄씩 차례로 입력하세요:')
doc.add_paragraph()
doc.add_paragraph('  1) Git 저장소 시작하기:')
add_code_block('git init')
doc.add_paragraph('  >> "이 폴더를 Git으로 관리할게!"라는 뜻')
doc.add_paragraph()
doc.add_paragraph('  2) 모든 파일을 올릴 준비하기:')
add_code_block('git add .')
doc.add_paragraph('  >> "이 폴더의 모든 파일을 포장해줘"라는 뜻 (마지막에 점(.) 꼭 입력!)')
doc.add_paragraph()
doc.add_paragraph('  3) 저장(커밋)하기:')
add_code_block('git commit -m "My first ETF dashboard project"')
doc.add_paragraph('  >> "이 상태를 기억해줘, 이름은 My first ETF dashboard project야"라는 뜻')
doc.add_paragraph()
doc.add_paragraph('  4) GitHub 저장소와 연결하기:')
doc.add_paragraph('  아래 명령어에서 "내깃허브이름"을 본인의 GitHub 사용자 이름으로 바꾸세요:')
add_code_block('git remote add origin https://github.com/내깃허브이름/etf-youtube-dashboard.git')
doc.add_paragraph()
doc.add_paragraph('  5) GitHub에 올리기!:')
add_code_block('git branch -M main\ngit push -u origin main')
doc.add_paragraph()
doc.add_paragraph('GitHub 로그인 창이 나타날 수 있어요:')
doc.add_paragraph('  - 브라우저가 열리면 "Authorize" 또는 "허용"을 클릭하세요')
doc.add_paragraph('  - 또는 GitHub 사용자 이름과 비밀번호를 입력하세요')
doc.add_paragraph()
doc.add_paragraph('성공하면 이런 메시지가 나와요:')
add_code_block('To https://github.com/내깃허브이름/etf-youtube-dashboard.git\n * [new branch]      main -> main')
doc.add_paragraph()

add_para('Step 4. GitHub에서 내 프로젝트 확인하기!', bold=True, size=12)
doc.add_paragraph('브라우저에서 아래 주소로 이동하세요 (내깃허브이름을 본인 이름으로):')
add_code_block('https://github.com/내깃허브이름/etf-youtube-dashboard')
doc.add_paragraph()
doc.add_paragraph('내가 만든 파일들이 GitHub 웹사이트에 예쁘게 보여요!')
doc.add_paragraph('이 주소를 친구에게 공유하면, 친구도 내 프로젝트를 볼 수 있습니다!')
doc.add_paragraph()

add_heading_styled('CHECK 체크포인트', level=3)
add_checkpoint([
    'CHECK GitHub에 가입했나요?',
    'CHECK etf-youtube-dashboard 저장소를 만들었나요?',
    'CHECK git push 후 GitHub 웹사이트에서 파일이 보이나요?',
])
doc.add_paragraph()

add_heading_styled('!! 막혔을 때', level=3)
add_error_table([
    ('git push에서 "Authentication failed"', '로그인 실패', '브라우저에서 GitHub에 로그인되어 있는지 확인. git push를 다시 실행하면 로그인 창이 뜰 수 있어요'),
    ('"remote origin already exists"', '이미 연결됨', 'git remote remove origin 을 먼저 실행하고 다시 git remote add... 하세요'),
    ('"Repository not found"', '저장소 이름이 다름', 'GitHub 웹사이트에서 저장소 이름을 확인하고, 명령어의 URL을 정확히 맞추세요'),
    ('아무것도 올라가지 않아요', 'add와 commit을 안 함', 'git add . -> git commit -m "메시지" -> git push 순서대로 다시 실행'),
])

doc.add_paragraph()
doc.add_paragraph()
add_para('4회차 & 전체 강의 완료! 축하합니다!!!', bold=True, size=14, color=(0, 128, 0))

doc.add_page_break()

# ════════════════════════════════════════════════════
# 전체 로드맵 요약표
# ════════════════════════════════════════════════════
add_heading_styled('전체 로드맵 요약표', level=1)
doc.add_paragraph()
add_para('4회차 강의 전체 흐름을 한눈에 확인하세요!', italic=True, size=12)
doc.add_paragraph()

table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = table.rows[0].cells
headers[0].text = '회차'
headers[1].text = '주제'
headers[2].text = '챕터'
headers[3].text = '핵심 결과물'
headers[4].text = '사용 도구'

roadmap_data = [
    ('1회차\n(30분)', '개발 환경 설정', '1. 터미널 설치\n2. Git 설치\n3. Node.js/Python 설치',
     '터미널, Git Bash 사용 가능\nnode, npm, python 작동', 'MS Store, Git,\nNode.js, Python'),
    ('2회차\n(30분)', 'Claude Code +\nAPI 키 발급', '4. Claude Code 설치\n5. YouTube API 키 발급',
     'AI 비서와 대화 가능\nAPI 키 + .env 파일 저장', 'npm, Claude Code,\nGoogle Cloud'),
    ('3회차\n(30분)', '데이터 수집\n자동화', '6. 수집 프로그램 생성\n7. 수집 실행',
     'etf_youtube_data.json\n(30개 영상 데이터)', 'Claude Code,\nPython'),
    ('4회차\n(30분)', '대시보드 제작\n+ GitHub 공유', '8. 대시보드 만들기\n9. 브라우저 확인 + 꾸미기\n보너스. GitHub 올리기',
     'index.html 대시보드 완성!\nGitHub에 프로젝트 공유', 'Claude Code,\n브라우저, GitHub'),
]

for row_data in roadmap_data:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()
add_para('총 소요 시간: 4회차 x 30분 = 약 2시간', bold=True, size=13)
doc.add_paragraph()

# 전체 흐름
add_heading_styled('전체 흐름 다이어그램', level=2)
doc.add_paragraph()
add_para(
    '[1회차] 터미널 설치 -> Git 설치 -> Node.js/Python 설치\n'
    '                            |\n'
    '[2회차] Claude Code 설치 -> API 키 발급 및 저장\n'
    '                            |\n'
    '[3회차] Claude에게 데이터 수집 프로그램 부탁 -> 실행\n'
    '                            |\n'
    '[4회차] Claude에게 대시보드 제작 부탁 -> 브라우저에서 확인!',
    size=11
)

doc.add_paragraph()
doc.add_paragraph()
add_heading_styled('마무리 - 여기까지 완주하신 여러분에게', level=2)
doc.add_paragraph()
doc.add_paragraph('축하합니다! 4회에 걸쳐 여러분은:')
doc.add_paragraph('  1. 터미널(까만 화면)을 두려움 없이 사용하게 되었고')
doc.add_paragraph('  2. Git이라는 개발자 필수 도구를 설치했고')
doc.add_paragraph('  3. Node.js와 Python이라는 프로그래밍 도구를 설치했고')
doc.add_paragraph('  4. AI(Claude Code)와 대화하며 프로그램을 만들었고')
doc.add_paragraph('  5. 유튜브 API로 실제 데이터를 수집했고')
doc.add_paragraph('  6. 그 데이터를 예쁜 대시보드로 시각화했고')
doc.add_paragraph('  7. GitHub에 내 프로젝트를 올려서 세상과 공유했습니다!')
doc.add_paragraph()
doc.add_paragraph('이 경험을 바탕으로 더 다양한 프로젝트에 도전해보세요.')
doc.add_paragraph('Claude Code에게 부탁하면 거의 모든 것을 만들 수 있습니다!')
doc.add_paragraph()
add_para('다음에 도전해볼 것들:', bold=True)
doc.add_paragraph('  - 다른 키워드(예: "비트코인", "부동산")로 데이터 수집해보기')
doc.add_paragraph('  - 대시보드에 검색 기능 추가해보기')
doc.add_paragraph('  - 매일 자동으로 데이터를 모아오게 만들기')
doc.add_paragraph('  - 수집 데이터를 엑셀 파일로 저장하기')
doc.add_paragraph('  - Git으로 내 코드를 GitHub에 공유해보기')

# ── 저장 ──
desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
filepath = os.path.join(desktop, 'ETF_유튜브_대시보드_강의안.docx')
doc.save(filepath)
print(f'saved: {filepath}')
